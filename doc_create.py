# ==========================формирование документа из шаблона===========================================================
# значения переменных читаются из БД PostgreSQL по шаблону и пишутся в файл <file_name><дата время>
import os
from sys import argv, exit
import datetime
import docx
import psycopg2
import openpyxl

try:
  prname, file_name,WorkDoc,WorkBPR,WorkBP,WorkENTP = argv
  varvel = [] # данные переменной из БД [переменная,значение,тип ]
  strnow = '' # текущий день и время
  vv =''
#==========================читаем из БД значение переменной=============================================================
  # def get_var_val(s,di,bpr,bp,entp):# считать значение переменной из БД s - переменная
  #    ss = f"'{s}'"
  #    sa = 'select docp_p,docp_v,docp_t,docp_r,docp_c from docp where docp_p='+ss+' and doc_id ='+di+' and bpr_id ='+bpr+ \
  #         ' and bp_id =' + bp +' and entp_id ='+entp
  #    cursor.execute(sa)
  #    return cursor.fetchone()
  def get_var_val(s):# считать атрибуты переменной из БД s - имя переменной
     global vv
     ss = f"'{s}'"
     wd = str(WorkDoc)
     wr = str(WorkBPR)
     wp = str(WorkBP)
     we = str(WorkENTP)
     if s[2] == 'E':
         wd = '0'
         wp = '0'
     elif s[2] == 'P':
         wd = '0'
     elif s[2] == 'R':
         wd = '0'
     sa = 'select docp_p,docp_v,docp_t,docp_r,docp_c,docp_id from docp where docp_p='+ss+ ' and doc_id ='+wd+\
          ' and bp_id =' + wp + ' and entp_id ='+we
     cursor.execute(sa)
     var=cursor.fetchone()
     sa = 'select docpv_v from docpv where docp_id=' + str(var[5]) + ' and doc_id =' + str(WorkDoc) + ' and bpr_id =' + wr
     cursor.execute(sa)
     vv = cursor.fetchone()
     return var
#============================находим переменную и ее значение===========================================================
  def var_rep(s):
      m = 0
      n = 0
      global varvel, cell, vv
      while m != -1:
          m = s.find('${', n)
          if m != -1:
              n = s.find('}', m + 1)
              if n == -1:
                print(' ошибка в шаблоне ${..')
                return  -1
              else:
                ss = s[m:n + 1]   # ключ
                wd = str(WorkDoc)
                wr = str(WorkBPR)
                wp = str(WorkBP)
                we = str(WorkENTP)
                if ss[2] == 'E':
                  wd = '0'
                  wr = '0'
                  wp = '0'
                elif ss[2] == 'P':
                  wd = '0'
                  wr = '0'
                elif ss[2] == 'R':
                  wd = '0'
                varvel = get_var_val(ss)
                if (varvel == None):
#                  print('значение переменной '+ss+' не задано')
                  return -1
                if varvel[2] == 'serial':
                  i = int(vv)
                  i += 1
                  vv = str(i)
                return 0
          else:
              return -1
#============================создание документа  по шаблону=============================================================
  def doc_cr(file_name):
      global vv
#------------------------------обработка файлов txt---------------------------------------------------------------------
      if ext[1] in ['.txt', '.html']:
        with open(file_name, 'r', encoding='UTF-8') as fi,\
             open(name, 'w', encoding='UTF-8') as fw:
          lines = fi.readlines()
          for line in lines:
              if len(line) > 4:
                if var_rep(line) == 0:
                   line = line.replace(varvel[0],varvel[1])
              fw.writelines(line)
          fw.close()
#-------------------------------python-docx----------------------------------------------------------------------------
      if ext[1] in ['.odf','.docx']:
         doc = docx.Document(file_name)
         if len(doc.paragraphs) > 1:
            for par in doc.paragraphs:
                if var_rep(par.text) == 0:
                   s = vv[0]
                   par.text = par.text.replace(varvel[0],s)
         n = len(doc.tables)
         if n > 0:
            for tab in doc.tables:
                for ro in tab.rows:
                   for cell in ro.cells:
                       if len(cell.text) > 3:
                          s =  cell.text
                          if var_rep(cell.text) == 0:
                             s = vv[0]
                             cell.text = cell.text.replace(varvel[0],s)
         doc.save(name)

#-------------------------------------xlrd------------------------------------------------------------------------------
      # if ext[1] =='.xls':
      #    workbook = xlrd.open_workbook(file_name)
      #    sheets_name = workbook.sheet_names()
      #    for names in sheets_name:
      #        worksheet = workbook.sheet_by_name(names)
      #        num_rows = worksheet.nrows
      #        num_cells = worksheet.ncols
      #        for row in range(num_rows):
      #            for sel in range(num_cells):
      #                val = worksheet.cell_value(row, sel)
      #                if val != None:
      #                  if (len(val) > 3):
      #                     var_rep(val)
      #        workbook.save(name)
#------------------------------------openpyxl---------------------------------------------------------------------------
      if ext[1] =='.xlsx':
         workbook = openpyxl.load_workbook(file_name,data_only=True)
         sheet = workbook.active
         for row in range(sheet.max_row):
             row += 1
             for sel in range(sheet.max_column):
                 sel += 1
                 val = sheet.cell(row, sel).value
                 if val != None:
                   if (len(val) > 3):
                      if var_rep(val) == 0:
                         sheet.cell(row,sel).value = sheet.cell(row,sel).value.replace(varvel[0],varvel[1])
         workbook.save(name)
 #======================основная программа=============================================================================
  print('Идет запись данных в файл. Ждите...')
  conn = psycopg2.connect(host='localhost', database='BP', user='postgres', password='rfn15')
  # Получаем объект курсора для выполнения SQL-запросов
  cursor = conn.cursor()
  conn.autocommit = True
  ext = os.path.splitext(file_name)
  now = datetime.datetime.now()
  strnow = now.strftime(" %d-%m-%y %H %M")
  pat =os.path.dirname(file_name)
  nn = ext[0][file_name.rfind('\\')+1:]
  l = len(pat)
  name = pat[0:l-9]+'BPR_'+str(WorkBPR)+'\\'+nn+' '+strnow + ext[1]
  doc_cr(file_name)
  cursor.close()
  conn.close()
  exit(0)

except FileNotFoundError:
       print('file not found-' + file_name)
       exit(-1)
except psycopg2.Error:
       print ('ошибка БД')
       exit(-1)