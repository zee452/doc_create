# ========================= чтение данных из всех бланков БП и запись их в docb без имен ппеременных+==================
import os
from sys import argv, exit
import docx
import psycopg2
# import openpyxl
import re
import zlib

try:
  prname, dir_name = argv
  #=====================================================================================================================
  def find_files(dir_path):# находим все файлы в папке бланки
     f = []
     for (dirpath, dirnames, filenames) in os.walk(dir_path):
         f.extend(filenames)
         break
     return f
#=======================================================================================================================
#  запись в БД атрибуты переменной замены
  def write_blank(s, r, c,crc, doc): # s -описание t- имя переменной
      ss = f"'{s}'"
      docs = f"'{os.path.basename(doc)}'"
      sa = 'insert into docb (docb_d,docb_r,docb_c,docb_p,doc_name) values ('+ss+ ',' + str(r) + ',' + \
            str(c)+','+str(crc)+','+docs+') on conflict do nothing'
      cursor.execute(sa)
 # =======================================================================================================================
 #  создание таблицы docp b docpv
  def docp_cr():  # s -описание t- имя переменной
      ss = f"'{s}'"
      docs = f"'{os.path.basename(doc)}'"
      sa = 'insert into docb (docb_d,docb_r,docb_c,docb_p,doc_name) values (' + ss + ',' + str(r) + ',' + \
           str(c) + ',' + str(crc) + ',' + docs + ') on conflict do nothing'
      cursor.execute(sa)


  #============================анализ бланка и создание переменной и ее атрибутов=========================================
  def var_cr(file_name):
#-------------------------------python-docx----------------------------------------------------------------------------
      if ext[1] in ['.odf','.docx']:
         doc = docx.Document(file_name)
         n = len(doc.tables)
         if n > 0:
            for tab in doc.tables:
                r = -1
                for ro in tab.rows:
                   r +=1
                   c = -1
                   pred = ''
                   for cell in ro.cells:
                       c += 1
                       s =  cell.text
                       if  s != pred and (re.search(r"[а-яА-Я]",s) != None):
                          if cell.paragraphs[0].runs[0].font.superscript:
                             ss =  s.lower().strip()
                             crc32 = zlib.crc32(ss.encode('utf-8'))
                             write_blank(s.lower().strip(), r, c,crc32, file_name) # пишем координаты переменной
                       pred = s
#-------------------------------------xlrd------------------------------------------------------------------------------
      # if ext[1] =='.xls':
      #    workbook = xlrd.open_workbook(file_name)
      #    sheets_name = workbook.sheet_names()
      #    for names in sheets_name:
      #        worksheet = workbook.sheet_by_name(names)
      #        num_rows = worksheet.nrows
      #        num_cells = worksheet.ncols
      #        for row in range(num_rows):
      #            for sel in range(num_cells):
      #                val = worksheet.cell_value(row, sel)
      #                if val != None:
      #                  if (len(val) > 3):
      #                     var_rep(val)
      #        workbook.save(name)
#------------------------------------openpyxl---------------------------------------------------------------------------
      # if ext[1] =='.xlsx':
      #    workbook = openpyxl.load_workbook(file_name,data_only=True)
      #    sheet = workbook.active
      #    for row in range(sheet.max_row):
      #        row += 1
      #        for sel in range(sheet.max_column):
      #            sel += 1
      #            val = sheet.cell(row, sel).value
      #            if val != None:
      #              if (len(val) > 3):
      #                 if var_rep(val) == 0:
      #                    sheet.cell(row,sel).value = sheet.cell(row,sel).value.replace(varvel[0],varvel[1])
      #    workbook.save(file_name)
 #======================основная программа=============================================================================
  print('Идет запись данных в БД. Ждите...')
  conn = psycopg2.connect(host='localhost', database='BP', user='postgres', password='rfn15')
  # Получаем объект курсора для выполнения SQL-запросов
  cursor = conn.cursor()
  conn.autocommit = True
  for file_name in find_files(dir_name):
      file_name = dir_name+'\\'+file_name
      ext = os.path.splitext(file_name)
      var_cr(file_name)
  docp_cr()
  cursor.close()
  conn.close()
  exit(0)

except FileNotFoundError:
       print('Path not found-' + dir_name)
       exit(-1)
except psycopg2.Error:
       print ('ошибка БД')
       exit(-1)