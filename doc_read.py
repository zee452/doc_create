# ==========================чтение документа файла, полученного из шаблона ()===========================================
# значения переменных читаются из файла и записываются в БД PostgreSQL
import os
from sys import argv, exit
import docx
import psycopg2
import openpyxl

try:
  prname, file_name,file_templ,WorkDoc,WorkBPR,WorkBP,WorkENTP = argv    # соответствующий файлу шаблон
  varvel = [] # данные переменной из БД [переменная,значение,тип, координаты и id  в таблице docp ]

#==========================читаем из БД значение переменной=============================================================

  def get_var_val(s):# считать атрибуты переменной из БД s - имя переменной
     ss = f"'{s}'"
     wd = str(WorkDoc)
     wp = str(WorkBP)
     we = str(WorkENTP)
     if s[2] == 'E':
         wd = '0'
         wp = '0'
     elif s[2] == 'P':
         wd = '0'
     elif s[2] == 'R':
         wd = '0'
     sa = 'select docp_p,docp_v,docp_t,docp_r,docp_c,docp_id from docp where docp_p='+ss+ ' and doc_id ='+wd+\
          ' and bp_id =' + wp + ' and entp_id ='+we
     cursor.execute(sa)
     return cursor.fetchone()
# ============================находим переменную ее значение и координаты в таблице======================================
  def var_find(s):  # s- поле шаблона
      m = 0
      n = 0
      global varvel, cell
      while m != -1:
          m = s.find('${', n)
          if m != -1:
              n = s.find('}', m + 1)
              if n == -1:
                  print(' ошибка в шаблоне ${..')
                  return -1
              else:
                  ss = s[m:n + 1]  # имя переменной
                  varvel = get_var_val(ss)  # находим ее атрибуты
                  if varvel != None:
                      return 0
                  else:
                      return -1
          else:
              return -1
  # ===================запись значения переменной в БД==================================================================
  def save_bd(s):           # записать значение переменной в БД
     ss = f"'{s}'"          # значение переменной docpv_v
     sss = f"'{varvel[5]}'" # docp_id
     wd = str(WorkDoc)
     wr = str(WorkBPR)

     if varvel[0][2] == 'E':
         sa = 'update docpv set docpv_v=' + ss + ' where docp_id=' + sss + ' and bpr_id <> 0'
     elif varvel[0][2] == 'P':
         sa = 'update docpv set docpv_v=' + ss + ' where docp_id=' + sss + ' and bpr_id <> 0'
     elif varvel[0][2] == 'R':
         sa = 'update docpv set docpv_v=' + ss + ' where docp_id=' + sss + ' and bpr_id =' + wr
     else:
         sa = 'update docpv set docpv_v=' + ss + ' where docp_id=' + sss + ' and bpr_id =' + wr + ' and doc_id ='+wd

     cursor.execute(sa)


#============================чтение  документа  по шаблону=============================================================
  def doc_read(file_name,file_templ):
#------------------------------обработка файлов txt---------------------------------------------------------------------
      if ext[1] in ['.txt', '.html']:
        with open(file_name, 'r', encoding='UTF-8') as fi,\
             open(file_templ, 'r', encoding='UTF-8') as fw:
          lines = fi.readlines()

#-------------------------------python-docx----------------------------------------------------------------------------
      if ext[1] in ['.odt','odf','odp','.docx']:
         doc = docx.Document(file_name)   # заполненный документ
         tmp = docx.Document(file_templ)  # шаблон документа
         if len(tmp.paragraphs) > 1:
            for par in tmp.paragraphs:
                if var_find(par.text) == 0:
                   print(' найти значение переменной невозможно')
                   return -1
         n = len(tmp.tables)
         if n > 0:
            n = -1
            for tab in tmp.tables:
                n += 1     # номер таблицы
                r = -1
                table = doc.tables[n]
                for ro in tab.rows:
                    r += 1
                    pred = ''
                    k = -1
                    for cell in ro.cells:
                        k +=1
                        # s = cell.text
                        # table = doc.tables[n]
                        # ss = table.cell(r - 1, k).text
                        if len(cell.text) > 3:
                           if var_find(cell.text) == 0: # находим переменную в шаблоне
                              if varvel[0] != pred:
                                s = table.cell(r,k).text
                                save_bd(s)       # значение переменной
                                pred = varvel[0]

#------------------------------------openpyxl---------------------------------------------------------------------------
      if ext[1] =='.xlsx':
         workbook = openpyxl.load_workbook(file_name,data_only=True)
         sheet = workbook.active
         for row in range(sheet.max_row):
             row += 1
             for sel in range(sheet.max_column):
                 sel += 1
                 val = sheet.cell(row, sel).value
                 if val != None:
                   if (len(val) > 3):
                      if var_find(val) == 0:
                         sheet.cell(row,sel).value = sheet.cell(row,sel).value.replace(varvel[0],varvel[1])

 #======================основная программа=============================================================================
  print('Идет запись данных в БД. Ждите...')
  conn = psycopg2.connect(host='localhost', database='BP', user='postgres', password='rfn15')
  # Получаем объект курсора для выполнения SQL-запросов
  cursor = conn.cursor()
  conn.autocommit = True
  ext = os.path.splitext(file_name)
  doc_read(file_name,file_templ)
  cursor.close()
  conn.close()
  exit(0)

except FileNotFoundError:
     print('file not found-' + file_name)
     exit(-1)